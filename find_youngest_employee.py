import os
from docx import Document
import re
from datetime import datetime

# Folder containing the application files
FOLDER = 'Scan3-20260415'

# Patterns to match birthdate (adjust as needed)
DATE_PATTERNS = [
    r'(\d{2}[./-]\d{2}[./-]\d{4})',  # 15.04.2000 or 15-04-2000
    r'(\d{4}[./-]\d{2}[./-]\d{2})',  # 2000-04-15
]

# Try to extract date from text
def extract_date(text):
    for pattern in DATE_PATTERNS:
        matches = re.finditer(pattern, text)
        for match in matches:
            date_str = match.group(1)
            # Try various formats
            for fmt in ("%d.%m.%Y", "%d-%m-%Y", "%d/%m/%Y", "%Y-%m-%d", "%Y.%m.%d", "%Y/%m/%d"):
                try:
                    return datetime.strptime(date_str, fmt)
                except ValueError:
                    continue
    return None

# Try to extract name from filename (as fallback)
def extract_name_from_filename(filename):
    name = os.path.splitext(filename)[0]
    return name.replace('_', ' ').replace('  ', ' ').strip()

# Try to extract name from document text
def extract_name_from_text(text):
    # Heuristic: look for capitalized words that look like a name (2 or 3 words)
    # This regex handles Turkmen and common Latin characters
    name_pattern = r'([A-ZÄÇĞIÄÖŞÜÝŽ][^\s]{2,}\s+[A-ZÄÇĞIÄÖŞÜÝŽ][^\s]{2,}(?:\s+[A-ZÄÇĞIÄÖŞÜÝŽ][^\s]{2,})?)'
    match = re.search(name_pattern, text)
    if match:
        return match.group(1).strip()
    return None

def main():
    youngest = None
    youngest_name = None
    youngest_file = None
    debug_count = 0
    for fname in os.listdir(FOLDER):
        if not fname.lower().endswith('.docx'):
            continue
        fpath = os.path.join(FOLDER, fname)
        try:
            doc = Document(fpath)
            # Extract text from paragraphs
            para_text = '\n'.join([p.text for p in doc.paragraphs])
            # Extract text from tables
            table_text = ''
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    table_text += row_text + '\n'
            full_text = para_text + '\n' + table_text
            # Print the text of the first 3 files for debugging
            if debug_count < 3:
                print(f"\n--- {fname} ---\n{full_text}\n--- END ---\n")
                debug_count += 1
            # Try to extract birthdate
            bdate = extract_date(full_text)
            # Try to extract name
            name = extract_name_from_text(full_text)
            if not name:
                name = extract_name_from_filename(fname)
            if bdate:
                if (youngest is None) or (bdate > youngest):
                    youngest = bdate
                    youngest_name = name
                    youngest_file = fname
        except Exception as e:
            print(f"Error reading {fname}: {e}")
    if youngest_name:
        print(f"Youngest employee: {youngest_name}")
        print(f"Birthdate: {youngest.strftime('%d.%m.%Y')}")
        print(f"File: {youngest_file}")
    else:
        print("No valid employee data found.")

if __name__ == "__main__":
    main()
